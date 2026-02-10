import hashlib
import io
import json
import os
import tempfile
import time
from datetime import date as _date
from typing import List, Optional

import fitz  # PyMuPDF
import streamlit as st
from docx import Document as DocxDocument
from PIL import Image
from google import genai
from google.genai import types

st.set_page_config(page_title="Bank Guarantee AI", layout="centered")

st.title("\U0001f3e6 Bank Guarantee AI")
st.write("Enter details strictly based on the Bank Guarantee document")

SHOW_STATUS = bool(st.secrets.get("SHOW_STATUS", False))


def _show_status(kind: str, message: str) -> None:
    if SHOW_STATUS:
        getattr(st, kind)(message)

MODELS_FALLBACK = [
    "gemini-2.5-flash-lite",
    "gemini-2.0-flash-lite",
    "gemini-2.5-flash",
    "gemini-3-flash-preview"
]
TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Data")
TEMPLATE_PERFORMANCE = os.path.join(TEMPLATE_DIR, "performance-bond-contract-letter.docx")
TEMPLATE_TENDER = os.path.join(TEMPLATE_DIR, "bid-bond-tender-letter.pdf.docx")
FIELD_KEYS = [
    "date",
    "bank_name",
    "guarantee_number",
    "guarantee_date",
    "amount",
    "company_name",
]


def _extract_json(text: str) -> Optional[dict]:
    text = text.strip()
    if text.startswith("{") and text.endswith("}"):
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return None
    try:
        return json.loads(text[start : end + 1])
    except json.JSONDecodeError:
        return None


def _pdf_to_images(pdf_bytes: bytes, max_pages: int = 5) -> List[Image.Image]:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    images: List[Image.Image] = []
    page_count = min(doc.page_count, max_pages)
    for page_index in range(page_count):
        page = doc.load_page(page_index)
        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        images.append(img)
    return images


def _pil_to_png_bytes(img: Image.Image) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def extract_fields_from_images(
    client: genai.Client,
    images: List[Image.Image],
    guarantee_type: str,
    model_name: str,
) -> Optional[dict]:
    prompt = (
        "You are extracting data from a bank guarantee document. "
        "The document may be in Arabic, English, or both. "
        "Only use text you can see in the document. "
        "For each field, provide TWO versions: one in English (translate if needed) and one in Arabic (translate if needed). "
        "Return JSON with exactly these keys: "
        "date, bank_name, bank_name_ar, guarantee_number, guarantee_date, amount, company_name, company_name_ar, guarantee_type, guarantee_type_ar. "
        "The '_ar' keys must contain the Arabic version. Non-_ar keys must be in English. "
        "If a field is missing, set it to null. "
        "Guarantee type is: " + guarantee_type + "."
    )

    contents: list = [prompt]
    for img in images:
        img_bytes = _pil_to_png_bytes(img)
        contents.append(
            types.Part.from_bytes(data=img_bytes, mime_type="image/png")
        )

    models_to_try = MODELS_FALLBACK if model_name is None else [model_name]
    last_exc = None

    for model_idx, current_model in enumerate(models_to_try):
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = client.models.generate_content(
                    model=current_model,
                    contents=contents,
                )
                return _extract_json(response.text)
            except Exception as exc:
                last_exc = exc
                err_str = str(exc)
                is_retryable = any(k in err_str for k in ("503", "429", "UNAVAILABLE", "RESOURCE_EXHAUSTED")) or "overloaded" in err_str.lower()

                if not is_retryable:
                    raise

                # Retry same model with exponential backoff
                if attempt < max_retries - 1:
                    wait = (attempt + 1) * 3
                    # st.warning(f"⏳ {current_model} — rate limited/busy, retrying in {wait}s (attempt {attempt + 1}/{max_retries})...")
                    time.sleep(wait)
                    continue

                # All retries exhausted → fall back to next model
                if model_idx < len(models_to_try) - 1:
                    next_model = models_to_try[model_idx + 1]
                    # st.warning(f"🔄 {current_model} failed after {max_retries} attempts, switching to {next_model}...")
                    break
                else:
                    st.error(f"All models exhausted. Last error: {err_str}")
                    raise

    return None


def build_letter_template(data: dict) -> str:
    guarantee_type = data.get("guarantee_type") or ""
    guarantee_type_ar = data.get("guarantee_type_ar") or guarantee_type
    date = data.get("date") or ""
    bank_name = data.get("bank_name") or ""
    bank_name_ar = data.get("bank_name_ar") or bank_name
    guarantee_number = data.get("guarantee_number") or ""
    guarantee_date = data.get("guarantee_date") or ""
    amount = data.get("amount") or ""
    company_name = data.get("company_name") or ""
    company_name_ar = data.get("company_name_ar") or company_name

    english = _build_english_letter(guarantee_type, date, bank_name, guarantee_number, guarantee_date, amount, company_name)
    arabic = _build_arabic_letter(guarantee_type_ar, date, bank_name_ar, guarantee_number, guarantee_date, amount, company_name_ar)

    return english + "\n" + ("=" * 60) + "\n\n" + arabic


def _build_english_letter(guarantee_type, date, bank_name, guarantee_number, guarantee_date, amount, company_name):
    header = (
        "Date: " + date + "\n"
        "To: The Chairman\n"
        "Municipality and Planning Department - Ajman\n"
        "P.O. Box 03 Ajman, UAE\n\n"
        "Subject: Bank Guarantee Confirmation\n\n"
        "Dear Sir,\n\n"
    )

    body = (
        "We hereby confirm the " + guarantee_type.lower() + " issued by " + bank_name + ". "
        "Guarantee No.: " + guarantee_number + ", dated " + guarantee_date + ", "
        "in the amount of " + amount + ", issued in favor of your department "
        "on behalf of " + company_name + ".\n\n"
    )

    footer = (
        "Should you require any further information, please contact us.\n\n"
        "Yours faithfully,\n"
        "For and on behalf of " + bank_name + "\n"
    )

    return header + body + footer


def _build_arabic_letter(guarantee_type, date, bank_name, guarantee_number, guarantee_date, amount, company_name):
    header = (
        "\u0627\u0644\u062a\u0627\u0631\u064a\u062e: " + date + "\n"
        "\u0625\u0644\u0649: \u0627\u0644\u0633\u064a\u062f \u0627\u0644\u0631\u0626\u064a\u0633\n"
        "\u0628\u0644\u062f\u064a\u0629 \u0648\u062f\u0627\u0626\u0631\u0629 \u0627\u0644\u062a\u062e\u0637\u064a\u0637 - \u0639\u062c\u0645\u0627\u0646\n"
        "\u0635.\u0628. 03 \u0639\u062c\u0645\u0627\u0646\u060c \u0627\u0644\u0625\u0645\u0627\u0631\u0627\u062a \u0627\u0644\u0639\u0631\u0628\u064a\u0629 \u0627\u0644\u0645\u062a\u062d\u062f\u0629\n\n"
        "\u0627\u0644\u0645\u0648\u0636\u0648\u0639: \u062a\u0623\u0643\u064a\u062f \u062e\u0637\u0627\u0628 \u0636\u0645\u0627\u0646 \u0628\u0646\u0643\u064a\n\n"
        "\u0633\u064a\u062f\u064a \u0627\u0644\u0639\u0632\u064a\u0632\u060c\n\n"
    )

    body = (
        "\u0646\u0624\u0643\u062f \u0628\u0645\u0648\u062c\u0628\u0647 " + guarantee_type + " \u0627\u0644\u0635\u0627\u062f\u0631 \u0645\u0646 " + bank_name + ". "
        "\u0631\u0642\u0645 \u0627\u0644\u0636\u0645\u0627\u0646: " + guarantee_number + "\u060c \u0628\u062a\u0627\u0631\u064a\u062e " + guarantee_date + "\u060c "
        "\u0628\u0645\u0628\u0644\u063a " + amount + "\u060c \u0644\u0635\u0627\u0644\u062d \u062f\u0627\u0626\u0631\u062a\u0643\u0645 "
        "\u0646\u064a\u0627\u0628\u0629 \u0639\u0646 " + company_name + ".\n\n"
    )

    footer = (
        "\u0644\u0623\u064a \u0627\u0633\u062a\u0641\u0633\u0627\u0631\u0627\u062a \u0625\u0636\u0627\u0641\u064a\u0629\u060c \u064a\u0631\u062c\u0649 \u0627\u0644\u062a\u0648\u0627\u0635\u0644 \u0645\u0639\u0646\u0627.\n\n"
        "\u0645\u0639 \u062e\u0627\u0644\u0635 \u0627\u0644\u062a\u0642\u062f\u064a\u0631\u060c\n"
        "\u0628\u0627\u0644\u0646\u064a\u0627\u0628\u0629 \u0639\u0646 " + bank_name + "\n"
    )

    return header + body + footer


# ---- Word Template Fill & PDF Conversion ----

def _set_cell_content(table, row_idx, col_idx, new_text):
    """Replace cell text while preserving the first run's formatting.
    Removes extra paragraphs so they don't add blank lines / push to page 2."""
    cell = table.cell(row_idx, col_idx)
    target_para = None
    for para in cell.paragraphs:
        if para.runs:
            target_para = para
            break
    if not target_para:
        return
    # Set new text in the first run, clear all other runs
    target_para.runs[0].text = new_text
    for run in target_para.runs[1:]:
        run.text = ""
    # Remove any extra paragraphs entirely from the cell XML
    from docx.oxml.ns import qn
    tc = cell._tc
    for para in list(cell.paragraphs):
        if para._element is not target_para._element:
            tc.remove(para._element)


def fill_word_template(data: dict, guarantee_type: str) -> bytes:
    """Open the correct Word template based on guarantee type, inject data, return .docx bytes."""
    if guarantee_type == "Performance Bond Guarantee":
        template_path = TEMPLATE_PERFORMANCE
    else:
        template_path = TEMPLATE_TENDER

    doc = DocxDocument(template_path)
    table = doc.tables[0]

    date = data.get("date", "")
    bank_name = data.get("bank_name", "")
    bank_name_ar = data.get("bank_name_ar", "") or bank_name
    guarantee_number = data.get("guarantee_number", "")
    guarantee_date = data.get("guarantee_date", "")
    amount = data.get("amount", "")
    company_name = data.get("company_name", "")
    company_name_ar = data.get("company_name_ar", "") or company_name

    # Row 0: Date (same for both templates)
    _set_cell_content(table, 0, 0, "\u0627\u0644\u062a\u0627\u0631\u064a\u062e: " + date)
    _set_cell_content(table, 0, 1, "Date: " + date)

    # Row 2: Bank name (same for both templates)
    _set_cell_content(table, 2, 0, "\u0627\u0644\u0633\u0627\u062f\u0629/ " + bank_name_ar)
    _set_cell_content(table, 2, 1, "M/s. " + bank_name)

    if guarantee_type == "Performance Bond Guarantee":
        # Performance Bond: rows 7, 8, 9, 10
        _set_cell_content(table, 7, 0, "\u0627\u0644\u0645\u0648\u0636\u0648\u0639: \u0636\u0645\u0627\u0646 \u0628\u0646\u0643\u064a \u0631\u0642\u0645: " + guarantee_number)
        _set_cell_content(table, 7, 1, "Subject: Bank Guarantee No. " + guarantee_number)
        _set_cell_content(table, 8, 0, "\u0628\u062a\u0627\u0631\u064a\u062e: " + guarantee_date)
        _set_cell_content(table, 8, 1, "Dated: " + guarantee_date)
        _set_cell_content(table, 9, 0, "\u0628\u0642\u064a\u0645\u0629: " + amount)
        _set_cell_content(table, 9, 1, "For AED " + amount)
        _set_cell_content(table, 10, 0, "\u0627\u0633\u0645 \u0627\u0644\u062c\u0647\u0629: " + company_name_ar)
        _set_cell_content(table, 10, 1, "M/s. " + company_name)
    else:
        # Tender Bond: rows 6, 7, 8, 9
        _set_cell_content(table, 6, 0, "\u0627\u0644\u0645\u0648\u0636\u0648\u0639: \u0636\u0645\u0627\u0646 \u0628\u0646\u0643\u064a \u0631\u0642\u0645: " + guarantee_number)
        _set_cell_content(table, 6, 1, "Subject: Bank Guarantee No. " + guarantee_number)
        _set_cell_content(table, 7, 0, "\u0628\u062a\u0627\u0631\u064a\u062e: " + guarantee_date)
        _set_cell_content(table, 7, 1, "Dated: " + guarantee_date)
        _set_cell_content(table, 8, 0, "\u0628\u0642\u064a\u0645\u0629: " + amount)
        _set_cell_content(table, 8, 1, "For AED " + amount)
        _set_cell_content(table, 9, 0, "\u0627\u0644\u0633\u0627\u062f\u0629/ " + company_name_ar)
        _set_cell_content(table, 9, 1, "M/s. " + company_name)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def convert_docx_to_pdf(docx_bytes: bytes) -> Optional[bytes]:
    """Convert .docx bytes to PDF using LibreOffice (works on Linux/Streamlit Cloud)."""
    import subprocess
    import shutil

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "letter.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                timeout=60,
                check=True,
                capture_output=True,
            )
        except Exception:
            return None
        pdf_path = os.path.join(tmpdir, "letter.pdf")
        if not os.path.exists(pdf_path):
            return None
        with open(pdf_path, "rb") as f:
            return f.read()


# --- Guarantee Type ---
guarantee_type = st.selectbox(
    "Select Guarantee Type",
    ["Tender Bond Guarantee", "Performance Bond Guarantee"],
)

st.divider()

# --- API key from secrets ---
api_key = st.secrets.get("GOOGLE_API_KEY", "")
if not api_key or api_key == "YOUR_API_KEY_HERE":
    st.error("Please set your GOOGLE_API_KEY in .streamlit/secrets.toml")
    st.stop()
auto_extract = True

# --- Upload File ---
uploaded_file = st.file_uploader(
    "Upload Bank Guarantee (PDF / Image)",
    type=["pdf", "png", "jpg", "jpeg"],
)

for key in FIELD_KEYS:
    st.session_state.setdefault(key, "")
st.session_state.setdefault("last_extracted_hash", None)

# Always set letter date to today
if not st.session_state["date"]:
    st.session_state["date"] = _date.today().strftime("%d/%m/%Y")

images: List[Image.Image] = []
if uploaded_file:
    _show_status("success", "File uploaded successfully")
    file_bytes = uploaded_file.read()
    file_hash = hashlib.sha256(file_bytes).hexdigest()
    if uploaded_file.type == "application/pdf":
        try:
            images = _pdf_to_images(file_bytes)
        except Exception as exc:
            st.error("PDF processing failed: " + str(exc))
    else:
        try:
            images = [Image.open(io.BytesIO(file_bytes))]
        except Exception as exc:
            st.error("Image processing failed: " + str(exc))

    if images:
        _show_status("caption", "Previewing " + str(len(images)) + " page(s)")
        st.image(images, width=400)

        should_auto_extract = (
            auto_extract
            and api_key
            and st.session_state["last_extracted_hash"] != file_hash
        )
        if should_auto_extract:
            client = genai.Client(api_key=api_key)
            with st.spinner("Extracting fields with AI..."):
                try:
                    result = extract_fields_from_images(
                        client, images, guarantee_type, None
                    )
                except Exception as exc:
                    result = None
                    st.error("Extraction failed: " + str(exc))

            if result:
                st.session_state["date"] = _date.today().strftime("%d/%m/%Y")
                st.session_state["bank_name"] = result.get("bank_name") or ""
                st.session_state["bank_name_ar"] = result.get("bank_name_ar") or ""
                st.session_state["guarantee_number"] = result.get("guarantee_number") or ""
                st.session_state["guarantee_date"] = result.get("guarantee_date") or ""
                st.session_state["amount"] = result.get("amount") or ""
                st.session_state["company_name"] = result.get("company_name") or ""
                st.session_state["company_name_ar"] = result.get("company_name_ar") or ""
                st.session_state["guarantee_type_ar"] = result.get("guarantee_type_ar") or ""
                st.session_state["last_extracted_hash"] = file_hash

st.divider()

# --- Data Fields ---
st.subheader("\U0001f4cb Guarantee Details")

date = st.text_input("Date", key="date")
bank_name = st.text_input("Bank Name", key="bank_name")
guarantee_no = st.text_input("Guarantee Number", key="guarantee_number")
guarantee_date = st.text_input("Date of Guarantee", key="guarantee_date")
amount = st.text_input("Amount", key="amount")
company_name = st.text_input("Company Name", key="company_name")

st.divider()

# --- Submit Button ---
cols = st.columns(2)
with cols[0]:
    extract_clicked = st.button("Extract With AI")
with cols[1]:
    generate_clicked = st.button("Generate Letter")

if extract_clicked:
    if not api_key:
        st.error("Please provide a Google API key")
    elif not images:
        st.error("Please upload a PDF or image first")
    else:
        client = genai.Client(api_key=api_key)
        with st.spinner("Extracting fields with AI..."):
            try:
                result = extract_fields_from_images(
                    client, images, guarantee_type, None
                )
            except Exception as exc:
                result = None
                st.error("Extraction failed: " + str(exc))

        if result:
            st.session_state["date"] = _date.today().strftime("%d/%m/%Y")
            st.session_state["bank_name"] = result.get("bank_name") or ""
            st.session_state["bank_name_ar"] = result.get("bank_name_ar") or ""
            st.session_state["guarantee_number"] = result.get("guarantee_number") or ""
            st.session_state["guarantee_date"] = result.get("guarantee_date") or ""
            st.session_state["amount"] = result.get("amount") or ""
            st.session_state["company_name"] = result.get("company_name") or ""
            st.session_state["company_name_ar"] = result.get("company_name_ar") or ""
            st.session_state["guarantee_type_ar"] = result.get("guarantee_type_ar") or ""

            _show_status("success", "Fields extracted. Please review and edit if needed.")
            if SHOW_STATUS:
                st.subheader("\U0001f4d1 Extracted Guarantee Information")
                st.write("**Guarantee Type:** " + guarantee_type)
                st.write("**Date:** " + st.session_state["date"])
                st.write("**Bank Name:** " + st.session_state["bank_name"])
                st.write("**Guarantee No:** " + st.session_state["guarantee_number"])
                st.write("**Date of Guarantee:** " + st.session_state["guarantee_date"])
                st.write("**Amount:** " + st.session_state["amount"])
                st.write("**Company Name:** " + st.session_state["company_name"])

if generate_clicked:
    if not all(
        [
            st.session_state["date"],
            st.session_state["bank_name"],
            st.session_state["guarantee_number"],
            st.session_state["guarantee_date"],
            st.session_state["amount"],
            st.session_state["company_name"],
        ]
    ):
        st.error("Please ensure all fields are filled")
    else:
        letter = build_letter_template(
            {
                "guarantee_type": guarantee_type,
                "guarantee_type_ar": st.session_state.get("guarantee_type_ar", ""),
                "date": st.session_state["date"],
                "bank_name": st.session_state["bank_name"],
                "bank_name_ar": st.session_state.get("bank_name_ar", ""),
                "guarantee_number": st.session_state["guarantee_number"],
                "guarantee_date": st.session_state["guarantee_date"],
                "amount": st.session_state["amount"],
                "company_name": st.session_state["company_name"],
                "company_name_ar": st.session_state.get("company_name_ar", ""),
            }
        )

        _show_status("success", "Letter generated successfully")

        # Split into English and Arabic parts
        parts = letter.split("=" * 60)
        if len(parts) == 2:
            st.subheader("\U0001f4dd English Letter")
            st.text_area("English", value=parts[0].strip(), height=280)
            st.subheader("\U0001f4dd Arabic Letter / \u0627\u0644\u062e\u0637\u0627\u0628 \u0628\u0627\u0644\u0639\u0631\u0628\u064a\u0629")
            st.text_area("Arabic", value=parts[1].strip(), height=280)
        else:
            st.subheader("\U0001f9fe Letter Preview")
            st.text_area("Letter", value=letter, height=320)

        # --- Generate Word & PDF from template ---
        st.divider()
        st.subheader("\U0001f4e5 Download Letter (Template)")

        template_data = {
            "date": st.session_state["date"],
            "bank_name": st.session_state["bank_name"],
            "bank_name_ar": st.session_state.get("bank_name_ar", ""),
            "guarantee_number": st.session_state["guarantee_number"],
            "guarantee_date": st.session_state["guarantee_date"],
            "amount": st.session_state["amount"],
            "company_name": st.session_state["company_name"],
            "company_name_ar": st.session_state.get("company_name_ar", ""),
        }

        active_template = TEMPLATE_PERFORMANCE if guarantee_type == "Performance Bond Guarantee" else TEMPLATE_TENDER
        if not os.path.exists(active_template):
            _show_status("warning", "Word template not found in Data folder.")
        else:
            with st.spinner("Generating Word document from template..."):
                docx_bytes = fill_word_template(template_data, guarantee_type)

            dl_cols = st.columns(2)

            with dl_cols[0]:
                st.download_button(
                    label="\U0001f4c4 Download Word (.docx)",
                    data=docx_bytes,
                    file_name="Bank_Guarantee_Letter.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            with dl_cols[1]:
                with st.spinner("Converting to PDF..."):
                    pdf_bytes = convert_docx_to_pdf(docx_bytes)
                if pdf_bytes:
                    st.download_button(
                        label="\U0001f4d5 Download PDF",
                        data=pdf_bytes,
                        file_name="Bank_Guarantee_Letter.pdf",
                        mime="application/pdf",
                    )
                else:
                    st.warning("PDF conversion failed. Please download the Word file instead.")
